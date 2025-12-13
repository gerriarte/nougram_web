"""
DOCX generation module for quotes
"""
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from typing import Optional


def format_currency(amount: float, currency: str = "USD") -> str:
    """Format currency amount"""
    currency_symbols = {
        "USD": "$",
        "COP": "$",
        "ARS": "$",
        "EUR": "€"
    }
    symbol = currency_symbols.get(currency, "$")
    
    # Format number with thousand separators
    if currency == "COP" or currency == "ARS":
        formatted = f"{amount:,.0f}"
    else:
        formatted = f"{amount:,.2f}"
    
    return f"{symbol} {formatted}"


def generate_quote_docx(
    project,
    quote,
    agency_name: str = "AgenciaOps"
) -> BytesIO:
    """
    Generate a professional DOCX quote
    
    Args:
        project: Project SQLAlchemy model instance
        quote: Quote SQLAlchemy model instance with items and project loaded
        agency_name: Agency name for header
    
    Returns:
        BytesIO buffer containing the DOCX
    """
    # Extract data from models
    project_name = project.name
    client_name = project.client_name
    client_email = project.client_email
    quote_version = quote.version
    currency = project.currency
    notes = quote.notes
    total_client_price = quote.total_client_price or 0
    margin_percentage = quote.margin_percentage or 0
    
    # Calculate taxes
    taxes = []
    total_taxes = 0
    if hasattr(project, 'taxes') and project.taxes:
        for tax in project.taxes:
            tax_amount = (total_client_price * tax.percentage / 100) if tax.percentage else 0
            taxes.append({
                'name': tax.name,
                'code': tax.code,
                'percentage': tax.percentage,
                'amount': tax_amount
            })
            total_taxes += tax_amount
    
    total_with_taxes = total_client_price + total_taxes
    
    # Extract quote items
    quote_items = []
    if hasattr(quote, 'items') and quote.items:
        for item in quote.items:
            service_name = item.service.name if item.service else 'Unknown Service'
            hours = item.estimated_hours or 0
            client_price = item.client_price or 0
            
            quote_items.append({
                'service_name': service_name,
                'estimated_hours': hours,
                'client_price': client_price
            })
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run(agency_name)
    header_run.font.size = Pt(24)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
    
    doc.add_paragraph("QUOTE", style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project and client info
    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    info_table.cell(0, 0).text = "Project:"
    info_table.cell(0, 1).text = project_name
    info_table.cell(1, 0).text = "Client:"
    info_table.cell(1, 1).text = client_name
    if client_email:
        info_table.cell(2, 0).text = "Email:"
        info_table.cell(2, 1).text = client_email
    info_table.cell(3, 0).text = "Quote Version:"
    info_table.cell(3, 1).text = f"v{quote_version}"
    
    # Style info table
    for row in info_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.bold = True if cell == row.cells[0] else False
    
    doc.add_paragraph()
    
    # Quote items table
    if quote_items:
        doc.add_paragraph("Services", style='Heading 2')
        
        items_table = doc.add_table(rows=1, cols=3)
        items_table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = items_table.rows[0].cells
        header_cells[0].text = "Service"
        header_cells[1].text = "Hours"
        header_cells[2].text = "Price"
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add items
        for item in quote_items:
            row_cells = items_table.add_row().cells
            row_cells[0].text = item['service_name']
            row_cells[1].text = f"{item['estimated_hours']:.1f}"
            row_cells[2].text = format_currency(item['client_price'], currency)
        
        doc.add_paragraph()
    
    # Totals
    doc.add_paragraph("Totals", style='Heading 2')
    
    totals_table = doc.add_table(rows=2 + len(taxes) + 1, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    row_idx = 0
    totals_table.cell(row_idx, 0).text = "Subtotal:"
    totals_table.cell(row_idx, 1).text = format_currency(total_client_price, currency)
    row_idx += 1
    
    # Taxes
    for tax in taxes:
        totals_table.cell(row_idx, 0).text = f"{tax['name']} ({tax['percentage']}%):"
        totals_table.cell(row_idx, 1).text = format_currency(tax['amount'], currency)
        row_idx += 1
    
    # Total
    totals_table.cell(row_idx, 0).text = "Total:"
    totals_table.cell(row_idx, 1).text = format_currency(total_with_taxes, currency)
    
    # Style totals - make total row bold
    total_row = totals_table.rows[row_idx]
    for cell in total_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Notes
    if notes:
        doc.add_paragraph("Notes", style='Heading 2')
        doc.add_paragraph(notes)
        doc.add_paragraph()
    
    # Footer
    footer_text = f"This quote is valid for 30 days from the date of issue. Margin: {(margin_percentage * 100):.1f}%"
    footer_para = doc.add_paragraph(footer_text)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(127, 140, 141)  # #7f8c8d
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


